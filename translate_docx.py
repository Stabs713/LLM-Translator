import re
import sys
from docx import Document
from docx.oxml.ns import qn
from tqdm import tqdm
from common import translate_chunk, chunk_text_by_sentences_safe

REFERENCE_TITLES = {
    'references', 'reference',
    'bibliography', 'bibliographie',
    'литература', 'список литературы', 'источники'
}


def extract_paragraph_with_math(paragraph):
    """
    Извлекает текст из параграфа, заменяя математические OMML блоки на placeholder'ы
    Возвращает: (текст_с_placeholder'ами, список_OMML_элементов)
    """
    math_elements = []
    parts = []

    for child in paragraph._element:
        # Обычный текст (run)
        if child.tag.endswith('r'):
            text = child.text or ""
            parts.append(text)

        # OMML математика (oMath)
        elif child.tag == qn('m:oMath') or child.tag == qn('m:oMathPara'):
            math_elements.append(child)
            placeholder = f"__MATH_{len(math_elements)-1}__"
            parts.append(placeholder)

    full_text = "".join(parts)
    return full_text, math_elements


def mask_text_formulas(text):
    """Маскирует текстовые формулы ($...$, $$...$$) если они есть"""
    placeholders = []

    def repl(match):
        placeholders.append(match.group(0))
        return f"__TEXTMATH_{len(placeholders)-1}__"

    # Display формулы
    text = re.sub(r'\$\$(.+?)\$\$', repl, text, flags=re.DOTALL)

    # Inline формулы
    text = re.sub(r'(?<!\$)\$(?!\$)([^$]+?)\$(?!\$)', repl, text)

    return text, placeholders


def unmask_text_formulas(text, placeholders):
    """Восстанавливает текстовые формулы"""
    for i in range(len(placeholders) - 1, -1, -1):
        text = text.replace(f"__TEXTMATH_{i}__", placeholders[i])
    return text


def rebuild_paragraph_with_math(paragraph, translated_text, math_elements):
    """
    Восстанавливает параграф с переведённым текстом и OMML элементами
    """
    # Очищаем параграф
    paragraph.clear()

    # Разбиваем переведённый текст по placeholder'ам математики
    # Используем regex чтобы точно захватить __MATH_число__
    parts = re.split(r'(__MATH_\d+__)', translated_text)

    for part in parts:
        # Проверяем, это placeholder математики?
        math_match = re.match(r'__MATH_(\d+)__', part)
        if math_match:
            # Это placeholder математики - вставляем OMML обратно
            idx = int(math_match.group(1))
            if idx < len(math_elements):
                paragraph._element.append(math_elements[idx])
        else:
            # Это обычный текст
            if part:
                paragraph.add_run(part)


def translate_docx(input_path, output_path):
    """Переводит DOCX файл с сохранением OMML формул"""
    try:
        doc = Document(input_path)
    except Exception as e:
        print(f"❌ Не удалось открыть .docx: {e}")
        sys.exit(1)

    in_references = False

    for para in tqdm(doc.paragraphs, desc="Перевод .docx"):
        if not para.text.strip():
            continue

        para_text_clean = para.text.strip().lower()
        if para_text_clean in REFERENCE_TITLES:
            in_references = True
            continue

        if in_references:
            continue

        # Извлекаем текст с placeholder'ами для OMML математики
        full_text, math_elements = extract_paragraph_with_math(para)

        if not full_text.strip():
            continue

        # Маскируем текстовые формулы (если есть $...$ в тексте)
        clean_text, text_formulas = mask_text_formulas(full_text)

        # Проверяем, есть ли что переводить
        temp_check = clean_text
        for i in range(len(math_elements)):
            temp_check = temp_check.replace(f"__MATH_{i}__", "")
        for i in range(len(text_formulas)):
            temp_check = temp_check.replace(f"__TEXTMATH_{i}__", "")

        if not re.search(r'[a-zA-Z]{2,}', temp_check):
            continue

        # Разбиваем на чанки и переводим
        chunks = chunk_text_by_sentences_safe(clean_text)
        translated = " ".join(
            translate_chunk(chunk) 
            for chunk in chunks if chunk.strip()
        )

        # Восстанавливаем текстовые формулы
        translated = unmask_text_formulas(translated, text_formulas)

        # Восстанавливаем параграф с OMML элементами
        rebuild_paragraph_with_math(para, translated, math_elements)

    doc.save(output_path)
    print(f"\n✅ Перевод завершён: {output_path}")
